import os
from docx import Document

# Function to read the text from a .docx file and process it
def extract_and_clean_sentences_from_file(file_path):
    # Open the .docx file using python-docx
    doc = Document(file_path)
    
    result = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        line = para.text
        
        # Check if the line is entirely uppercase
        if line.isupper():
            continue  # Skip this line if it's all uppercase
        
        # If not uppercase, add the line to the result
        result.append(line)
    
    return result

# Function to save the cleaned sentences to a new .docx file
def save_to_docx(output_path, cleaned_lines):
    doc = Document()
    for line in cleaned_lines:
        doc.add_paragraph(line)
    doc.save(output_path)

# Input file path and output file path
input_file_path = os.path.expanduser('~/Downloads/cleanv1.docx')
output_file_path = os.path.expanduser('~/Downloads/cleanv2.docx') 

# Check if input file exists
if not os.path.exists(input_file_path):
    print(f"File not found: {input_file_path}")
else:
    # Extract and clean sentences from the input file
    cleaned_lines = extract_and_clean_sentences_from_file(input_file_path)

    # Save the cleaned content to a new .docx file
    save_to_docx(output_file_path, cleaned_lines)

    print(f"Cleaned content has been saved to {output_file_path}")
